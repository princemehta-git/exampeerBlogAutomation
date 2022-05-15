import tkinter as tk
from tkinter.filedialog import askopenfilename, askdirectory
from tkinter import ttk
from datetime import datetime, timedelta
from PIL import ImageTk, Image
import glob
import subprocess
import winsound
# from gtts import gTTS
# # from playsound import playsound
# from pygame import mixer
import pyttsx3

import docx
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from googlesearch import search
import calendar
from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
import pyperclip
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
import time

main_window = tk.Tk()

main_window.geometry("800x500")
main_window.title('PostBlog')
main_window.iconbitmap('icon.ico')

currentDate = str(datetime.now().day)
currentMonth = str(calendar.month_abbr[datetime.now().month]).lower()
currentYear = str(datetime.now().year)

folder_path_var = tk.StringVar()
eng_filepath_var = tk.StringVar()
hin_filepath_var = tk.StringVar()
is_eng_var = tk.IntVar()
is_hin_var = tk.IntVar()
url_var = tk.IntVar()
wait_var = tk.IntVar()
is_showBrowser_var = tk.IntVar()
is_autoDetector_var = tk.IntVar()
is_autoAnalyser_var = tk.IntVar()
is_autoUploader_var = tk.IntVar()


bg = ImageTk.PhotoImage(Image.open('bg.png'))
tk.Label( main_window, image = bg).place(x = -2,y = -2)


folder_path_label = tk.Label(main_window, text='Folder Path', font=('Bradley Hand ITC', 20, 'bold'), bg = "#000000", fg = '#D1FF00')
folder_path_entry = tk.Entry(main_window, textvariable=folder_path_var, font=('Calibre', 10, 'italic'), width = 45)

eng_filepath_label = tk.Label(main_window, text='Eng Docx', font=('Bradley Hand ITC', 20, 'bold'), bg = "#000000", fg = '#D1FF00')
eng_filepath_entry = tk.Entry(main_window, textvariable=eng_filepath_var, font=('Calibre', 10, 'italic'), width = 45)

hin_filepath_label = tk.Label(main_window, text='Hin Docx', font=('Bradley Hand ITC', 20, 'bold'), bg = "#000000", fg = '#D1FF00')
hin_filepath_entry = tk.Entry(main_window, textvariable=hin_filepath_var, font=('Calibre', 10, 'italic'), width = 45)


url_label = tk.Label(main_window, text='Which url:', font=('Bradley Hand ITC', 16, 'bold'), bg = "#000000", fg = '#D1FF00')
url_label.place(relx=0.7, rely=0.8, relheight=0.05, relwidth=0.15)


which_url = ttk.Combobox(main_window, textvariable = url_var)
which_url.place(relx=0.85, rely=0.8, relheight=0.05, relwidth=0.06)

which_url['values'] = [1,2,3,4,5,6,7,8,9,10]

which_url.set(2)

wait_label = tk.Label(main_window, text='Wait Till(in secs)', font=('Bradley Hand ITC', 16, 'bold'), bg = "#000000", fg = '#D1FF00')
wait_entry = tk.Entry(main_window, textvariable=wait_var, font=('Calibre', 10, 'italic'), width = 45)

wait_label.place(relx=0.01, rely=0.8, relheight=0.035, relwidth=0.3)
wait_entry.place(relx=0.3, rely=0.8, relheight=0.035, relwidth=0.06)

wait_entry.delete(0, 'end')
wait_entry.insert(index=0,string=20)
wait = int(wait_var.get())




def browsefolder():
    filename = askdirectory()
    folder_path_entry.delete(0, 'end')
    folder_path_entry.insert(index=0,string=filename)  # add this
    try:
        docx_files = glob.glob(filename + "\\*.docx")
        # print(docx_files)
        for file in docx_files:
            if 'eng' in file.lower():
                is_eng_button.select()
                eng_filepath_entry.delete(0, 'end')
                eng_filepath_entry.insert(index=0, string=file)
                update()
            elif 'hin' in file.lower():
                is_hin_button.select()
                hin_filepath_entry.delete(0, 'end')
                hin_filepath_entry.insert(index=0, string=file)
                update()
    except:
        pass
    return None

def browse_eng_doc():
    filename = askopenfilename(filetypes=(("word | docx file", "*.docx"),("All files", " *.* "),))
    eng_filepath_entry.delete(0, 'end')
    eng_filepath_entry.insert(index=0,string=filename)  # add this
    return None

def browse_hin_doc():
    filename = askopenfilename(filetypes=(("word | docx file", "*.docx"),("All files", " *.* "),))
    hin_filepath_entry.delete(0, 'end')
    hin_filepath_entry.insert(index=0,string=filename)  # add this
    return None

brow_eng_btn = tk.Button(main_window, text="Browse", font=('Lucida Calligraphy', 10, 'bold'), command=browse_eng_doc,bg="#000000", fg='#FF00D8')
brow_hin_btn = tk.Button(main_window, text="Browse", font=('Lucida Calligraphy', 10, 'bold'), command=browse_hin_doc, bg = "#000000", fg = '#FF00D8')



folder_path_label.place(relx=0.01, rely=0.05, relwidth=0.2, relheight=0.035)
folder_path_entry.place(relx=0.2, rely=0.05, relwidth=0.68, relheight=0.035)
tk.Button(main_window, text="Browse", font=('Lucida Calligraphy', 10, 'bold'), command=browsefolder, bg = "#000000", fg = '#FF00D8').place(relx=0.9, rely=0.05, relwidth=0.08, relheight=0.035)


# eng_filepath_label.place(relx=0.01, rely=0.21, relwidth=0.2, relheight=0.035)
# eng_filepath_entry.place(relx=0.2, rely=0.21, relwidth=0.68, relheight=0.035)
# brow_eng_btn.place(relx=0.9, rely=0.21, relwidth=0.08, relheight=0.035)
#
# hin_filepath_label.place(relx=0.01, rely=0.29, relwidth=0.2, relheight=0.035)
# hin_filepath_entry.place(relx=0.2, rely=0.29, relwidth=0.68, relheight=0.035)
# brow_hin_btn.place(relx=0.9, rely=0.29, relwidth=0.08, relheight=0.035)

def update():
    if not is_eng_var.get():
        eng_filepath_label.place_forget()
        eng_filepath_entry.place_forget()
        brow_eng_btn.place_forget()
    else:
        eng_filepath_label.place(relx=0.01, rely=0.21, relwidth=0.2, relheight=0.035)
        eng_filepath_entry.place(relx=0.2, rely=0.21, relwidth=0.68, relheight=0.035)
        brow_eng_btn.place(relx=0.9, rely=0.21, relwidth=0.08, relheight=0.035)
    if not is_hin_var.get():
        hin_filepath_label.place_forget()
        hin_filepath_entry.place_forget()
        brow_hin_btn.place_forget()
    else:
        hin_filepath_label.place(relx=0.01, rely=0.29, relwidth=0.2, relheight=0.035)
        hin_filepath_entry.place(relx=0.2, rely=0.29, relwidth=0.68, relheight=0.035)
        brow_hin_btn.place(relx=0.9, rely=0.29, relwidth=0.08, relheight=0.035)




is_eng_button = tk.Checkbutton(main_window, text = "Eng file",
                      variable = is_eng_var,
                      onvalue = 1,
                      offvalue = 0,
                      height = 2,
                      width = 10,font=('MV Boli', 12, 'normal'), bg = "#000000", fg = '#FF00D8', command = update)

is_eng_button.place(relx=0.4, rely=0.13, relheight=0.05, relwidth=0.1)

is_hin_button = tk.Checkbutton(main_window, text = "Hin file",
                      variable = is_hin_var,
                      onvalue = 1,
                      offvalue = 0,
                      height = 2,
                      width = 10,font=('MV Boli', 12, 'normal'), bg = "#000000", fg = '#FF00D8', command = update)

is_hin_button.place(relx=0.6, rely=0.13, relheight=0.05, relwidth=0.1)

def autoDetector():
    
    with open('autoDetector.txt','w', encoding="utf-8") as file:
        file.write(str(is_autoDetector_var.get()))
    if is_autoDetector_var.get():
        # print(1)
        import os
        username = os.getenv('username')
        folderPath = 'C:\\Users\\%s\\Dropbox\\CURRENT AFFAIRS\\'%username
        # print(folderPath)
        for folder in os.listdir(folderPath):
            if currentMonth.lower() in folder.lower() and currentYear[-2:].lower() in folder.lower():
                folderPath += folder+'\\'
                break
        for folder in os.listdir(folderPath):
            if 'daily' in folder.lower() and 'update' in folder.lower():
                folderPath += folder + '\\'
                break
        for folder in os.listdir(folderPath):
            if 'daily' in folder.lower() and 'update' in folder.lower():
                folderPath += folder + '\\'
                break
        for folder in os.listdir(folderPath):
            if currentDate.lower() in folder.lower() and currentMonth.lower() in folder.lower():
                folderPath += folder
                break
        # print(folderPath)
        folder_path_entry.delete(0, 'end')
        folder_path_entry.insert(index=0, string=folderPath)
        docx_files = glob.glob(folderPath + "\\*.docx")
        # print(docx_files)
        for file in docx_files:
            if 'eng' in file.lower():
                is_eng_button.select()
                eng_filepath_entry.delete(0, 'end')
                eng_filepath_entry.insert(index=0,string=file)
                update()
            elif 'hin' in file.lower():
                is_hin_button.select()
                hin_filepath_entry.delete(0,'end')
                hin_filepath_entry.insert(index=0,string=file)
                update()
    return

def autoAnalyser():
    with open('autoAnalyser.txt','w', encoding="utf-8") as file:
        file.write(str(is_autoAnalyser_var.get()))

    if is_autoAnalyser_var.get():
        if not is_autoDetector_var.get():
            is_autoDetector_button.select()
            autoDetector()
        analyse()

    return

def autoUploader():
    global analyse_error
    with open('autoUploader.txt','w', encoding="utf-8") as file:
        file.write(str(is_autoUploader_var.get()))

    if is_autoUploader_var.get():
        if not is_autoAnalyser_var.get():
            is_autoAnalyser_button.select()
            autoAnalyser()
        if analyse_error == 0:
            upload()
        else:
            popup('Error in analysis\nso didnt uploaded blog', 'analysed.xlsx', 'error')
            return 



is_autoDetector_button = tk.Checkbutton(main_window, text = "AutoDetect",
                      variable = is_autoDetector_var,
                      onvalue = 1,
                      offvalue = 0,
                      height = 2,
                      width = 10,font=('MV Boli', 12, 'normal'), bg = "#000000", fg = '#FF00D8', command = autoDetector)

is_autoDetector_button.place(relx=0.02, rely=0.9, relheight=0.05, relwidth=0.2)

is_autoAnalyser_button = tk.Checkbutton(main_window, text = "AutoAnalyser",
                      variable = is_autoAnalyser_var,
                      onvalue = 1,
                      offvalue = 0,
                      height = 2,
                      width = 10,font=('MV Boli', 12, 'normal'), bg = "#000000", fg = '#FF00D8', command = autoAnalyser)

is_autoAnalyser_button.place(relx=0.4, rely=0.9, relheight=0.05, relwidth=0.2)

is_autoUploader_button = tk.Checkbutton(main_window, text = "AutoUploader",
                      variable = is_autoUploader_var,
                      onvalue = 1,
                      offvalue = 0,
                      height = 2,
                      width = 10,font=('MV Boli', 12, 'normal'), bg = "#000000", fg = '#FF00D8', command = autoUploader)

is_autoUploader_button.place(relx=0.8, rely=0.9, relheight=0.05, relwidth=0.2)

options = webdriver.ChromeOptions()
def showBrowser():
    # global options
    with open('showBrowser.txt','w', encoding="utf-8") as file:
        file.write(str(is_showBrowser_var.get()))
    # options.headless = not bool(is_showBrowser_var.get())

is_showBrowser_button = tk.Checkbutton(main_window, text = "ShowBrowser",
                      variable = is_showBrowser_var,
                      onvalue = 1,
                      offvalue = 0,
                      height = 2,
                      width = 10,font=('MV Boli', 12, 'normal'), bg = "#000000", fg = '#FF00D8', command = showBrowser)
is_showBrowser_button.place(relx=0.4, rely=0.6, relwidth=0.2, relheight=0.05)


def popup(txt, path, error_or_notification):
    global bg1, bg2
    top = tk.Toplevel(main_window)

    if 'err' in error_or_notification.lower():
        top.title('ERROR')
        top.iconbitmap('error.ico')
        bg1 = ImageTk.PhotoImage(Image.open('error.jpg'))
        tk.Label(top, image=bg1).place(x=-2, y=-2)
        # winsound.MessageBeep(winsound.MB_ICONHAND)


    else:
        top.title('PostBlog notifier')
        top.iconbitmap('icon.ico')
        bg2 = ImageTk.PhotoImage(Image.open('notification.jpg'))
        tk.Label(top, image=bg2).place(x=-2, y=-2)
        # winsound.MessageBeep(winsound.MB_OK)
    top.geometry("400x200")


    # Create an Entry Widget in the Toplevel window
    tk.Label(top, text=txt, font=('Bradley Hand ITC', 20, 'bold'), bg = "#000000", fg = '#D1FF00').pack()

    # btn for open folder
    tk.Button(top, text = 'Open',font=('Lucida Calligraphy', 10, 'bold'),bg = "#000000", fg = '#FF00D8', command=lambda:subprocess.call('explorer %s'%path, shell=True)).pack(pady=5, side=tk.TOP)


    # Create a Button Widget in the Toplevel Window
    button = tk.Button(top, text="Cancel",font=('Lucida Calligraphy', 10, 'bold'),bg = "#000000", fg = '#FF00D8', command=lambda :top.destroy())
    button.pack(pady=5, side=tk.TOP)
    try:

        engine = pyttsx3.init()
        voices = engine.getProperty('voices')
        engine.setProperty('voice', voices[1].id)
        engine.setProperty("rate", 230)
        engine.say(txt)
        engine.runAndWait()
    except:
        try:
            engine = pyttsx3.init()
            voices = engine.getProperty('voices')
            engine.setProperty('voice', voices[1].id)
            engine.setProperty("rate", 230)
            engine.say(txt)
            engine.runAndWait()
        except:
            pass



    return None


def analyse():
    global analyse_error
    analyse_error = 0
    path = folder_path_var.get()
    img_files = glob.glob(path + "\*.jpg")
    if len(img_files) != 10:
        rest_img_files = glob.glob(path + '\*.png') + glob.glob(path + '\*.svg') + glob.glob(
            path + '\*.webp') + glob.glob(path + '\*.jpeg') + glob.glob(path + '\*.jfif') + glob.glob(
            path + '\*.pjpeg') + glob.glob(path + '\*.pjp')
        for img in rest_img_files:
            Image.open(img).convert("RGB").save(img.replace(img.split('.')[-1], 'jpg'))

    img_files = glob.glob(path + "\*.jpg")
    if len(img_files) != 10:
        popup("Folder doesn't have \n 10 jpg images", path, 'error')
        analyse_error += 1
        return

    docx_files = glob.glob(path + "\*.docx")
    if len(docx_files) != is_hin_var.get() + is_eng_var.get():
        popup("Folder doesn't have \n %s docx files" % is_hin_var.get() + is_eng_var.get(), path, 'error')
        analyse_error += 1
        return

    wb = load_workbook('sample.xlsx')
    sheet = wb[wb.sheetnames[0]]

    if is_eng_var.get():
        eng_paras = docx.Document(eng_filepath_var.get()).paragraphs
        topics = ['Agreements', 'Appointments', 'Awards', 'Banking', 'Books and Authors', 'Defence', 'Economy',
                  'Important Days', 'International', 'Miscellaneous', 'National', 'Obituaries', 'Ranks & Reports',
                  'Schemes and Committees', 'Science & Technology', 'Sports', 'States', 'Summits and Conference']

        row_count = 1
        for para_count, para in enumerate(eng_paras):
            doctopic = ''
            if 'daily' in para.text.lower() and 'update' in para.text.lower() and len(para.text) < 65:
                row_count += 1

                # adding doctopic
                doctopic = para.text.split(':')[-1].strip()
                sheet.cell(row_count, 1).value = doctopic

                # adding matched_topic
                matched_topic = ''
                if doctopic.lower().strip() == 'national':
                    matched_topic = 'National'
                else:
                    for matched_topic in topics:
                        if doctopic[:4].lower() in matched_topic.lower():
                            break
                        else:
                            matched_topic = 'Miscellaneous'
                sheet.cell(row_count, 2).value = matched_topic

                # adding engtitle
                engtitle = ''
                track_tit = None
                if ((eng_paras[para_count + 1].text.strip() != '') and (
                        eng_paras[para_count + 1].text.strip() != None)):
                    engtitle = eng_paras[para_count + 1].text.strip().strip(':-')
                    track_tit = 1


                elif ((eng_paras[para_count + 2].text.strip() != '') and (
                        eng_paras[para_count + 2].text.strip() != None)):
                    engtitle = eng_paras[para_count + 2].text.strip().strip(':-')
                    track_tit = 2

                else:
                    analyse_error += 1
                    sheet.cell(row_count, 3).fill = PatternFill('solid', fgColor='FF0000')
                sheet.cell(row_count, 3).value = engtitle

                # adding engdiscription
                engdis = ''
                if ((eng_paras[para_count + 2].text.strip() != '') and (
                        eng_paras[para_count + 2].text.strip() != None)) and track_tit != 2:
                    engdis = eng_paras[para_count + 2].text.strip()

                elif ((eng_paras[para_count + 3].text.strip() != '') and (
                        eng_paras[para_count + 3].text.strip() != None)):
                    engdis = eng_paras[para_count + 3].text.strip()

                elif ((eng_paras[para_count + 4].text.strip() != '') and (
                        eng_paras[para_count + 4].text.strip() != None)):
                    engdis = eng_paras[para_count + 4].text.strip()

                else:
                    analyse_error += 1
                    sheet.cell(row_count, 4).fill = PatternFill('solid', fgColor='FF0000')
                sheet.cell(row_count, 4).value = engdis

                # adding url
                link = ''
                try:
                    link = search(engtitle, tld="co.in", num=1, start=int(url_var.get()) - 1, stop=1, pause=1.5,
                                  lang='english').__next__()
                except:
                    analyse_error += 1
                    sheet.cell(row_count, 5).fill = PatternFill('solid', fgColor='FF0000')

                sheet.cell(row_count, 5).value = link

                # adding img:
                img = ''
                for img in img_files:
                    if row_count - 1 == 1:
                        if '1'.lower() in img.split('\\')[-1].lower() and '10'.lower() not in img.split('\\')[-1].lower():
                            break
                    elif str(row_count - 1).lower() in img.split('\\')[-1].lower():
                        break
                if img == '':
                    analyse_error += 1
                    sheet.cell(row_count, 10).fill = PatternFill('solid', fgColor='FF0000')

                sheet.cell(row_count, 10).value = img
        wb.save('analysed.xlsx')

        if sheet.max_row != 11:
            popup('It couldnt find 10 topics of eng', 'analysed.xlsx', 'error')
            analyse_error += 1
            # return

    else:
        popup('Must Need Eng docx file \n to track titles', folder_path_var.get(), 'error')
        analyse_error += 1
        return

        # hindi file
    if is_hin_var.get():
        hin_paras = docx.Document(hin_filepath_var.get()).paragraphs

        row_count = 1
        for para_count, para in enumerate(hin_paras):
            if 'डेली' in para.text.lower() and 'अपडेट' in para.text.lower() and len(para.text) < 65:
                row_count += 1

                # adding hintitle
                hintitle = ''
                track_tit = None
                if ((hin_paras[para_count + 1].text.strip() != '') and (
                        hin_paras[para_count + 1].text.strip() != None)):
                    hintitle = hin_paras[para_count + 1].text.strip().strip(':-')
                    track_tit = 1


                elif ((hin_paras[para_count + 2].text.strip() != '') and (
                        hin_paras[para_count + 2].text.strip() != None)):
                    hintitle = hin_paras[para_count + 2].text.strip().strip(':-')
                    track_tit = 2
                else:
                    analyse_error += 1
                    sheet.cell(row_count, 6).fill = PatternFill('solid', fgColor='FF0000')

                sheet.cell(row_count, 6).value = hintitle

                # adding hinslug
                try:
                    sheet.cell(row_count, 7).value = 'hindi-%s-%s-%s' % (currentDate, currentMonth, row_count - 1)
                except:
                    analyse_error += 1
                    sheet.cell(row_count, 7).fill = PatternFill('solid', fgColor='FF0000')

                # adding hindiscription
                hindis = ''
                if ((hin_paras[para_count + 2].text.strip() != '') and (
                        hin_paras[para_count + 2].text.strip() != None)) and track_tit != 2:
                    hindis = hin_paras[para_count + 2].text.strip()

                elif ((hin_paras[para_count + 3].text.strip() != '') and (
                        hin_paras[para_count + 3].text.strip() != None)):
                    hindis = hin_paras[para_count + 3].text.strip()

                elif ((hin_paras[para_count + 4].text.strip() != '') and (
                        hin_paras[para_count + 4].text.strip() != None)):
                    hindis = hin_paras[para_count + 4].text.strip()

                else:
                    analyse_error += 1
                    sheet.cell(row_count, 8).fill = PatternFill('solid', fgColor='FF0000')

                sheet.cell(row_count, 8).value = hindis

                # adding url
                link = ''
                try:
                    link = search(hintitle, tld="co.in", num=1, start=int(url_var.get()) - 1, stop=1, pause=1.5,
                                  lang='hindi').__next__()
                except:
                    analyse_error += 1
                    sheet.cell(row_count, 9).fill = PatternFill('solid', fgColor='FF0000')

                sheet.cell(row_count, 9).value = link
        wb.save('analysed.xlsx')

        if sheet.max_row != 11:
            popup('It couldnt find 10 topics of eng', 'analysed.xlsx', 'error')
            analyse_error += 1
            return

    wb.save('analysed.xlsx')

    popup('Analysis Complete!!\n Error: %s' % analyse_error, 'analysed.xlsx', 'notification')
    return

analyse_btn = tk.Button(main_window, text='Analyse',font=('Lucida Calligraphy', 10, 'bold'),bg = "#000000", fg = '#FF00D8', command=analyse)
analyse_btn.place(relx=0.45, rely=0.4, relwidth=0.09, relheight=0.05)


def upload():
    global options
    error = 0
    sent = 0

    # options.headless = True
    user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/85.0.4183.83 Safari/537.36"
    options.add_argument(f'user-agent={user_agent}')
    options.add_argument("--window-size=1920,1080")
    options.add_argument('--ignore-certificate-errors')
    options.add_argument('--allow-running-insecure-content')
    options.add_argument("--disable-extensions")
    options.add_argument("--proxy-server='direct://'")
    options.add_argument("--proxy-bypass-list=*")
    options.add_argument("--start-maximized")
    options.add_argument('--disable-gpu')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--no-sandbox')
    options.add_argument("--app=https://www.google.com")
    options.add_experimental_option("useAutomationExtension", False)
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    options.headless = not bool(is_showBrowser_var.get())

    try:
        driver = webdriver.Chrome(ChromeDriverManager().install(), options=options)
        driver.implicitly_wait(wait)
        driver.maximize_window()
    except:
        popup('Update Chrome',r'C:\Program Files (x86)\Google\Chrome\Application\chrome.exe','error')
        return
    login = False
    while not login:
        driver.get('https://currentaffairsnews.in/admin-login')
        driver.find_element_by_id('input-email').send_keys('admin@gmail.com')
        driver.find_element_by_id('input-password').send_keys('admin')
        driver.find_element_by_id('btn-login').click()
        try:
            driver.find_elements_by_xpath('//a[@class="breadcrumb--active"]')
            login = True
        except:
            login = False


    wb1 = load_workbook('analysed.xlsx')
    sheet1 = wb1[wb1.sheetnames[0]]

    if is_eng_var.get():
        for row in range(2,12):
            sheet1.cell(row, 11).value = ''
            try:
                driver.get('https://currentaffairsnews.in/add-blog/side-menu/light')
            except:
                driver.get('https://currentaffairsnews.in/add-blog/side-menu/light')

            # img uploading
            img = sheet1.cell(row, 10).value

            try:
                driver.find_elements_by_xpath('//input[@id="image"]')[0].send_keys(img)
            except:
                try:
                    driver.find_elements_by_xpath('//input[@id="image"]')[0].send_keys(img)
                except:
                    error += 1
                    sheet1.cell(row, 10).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 11).value += ", Error while uploading img"
                    continue


            # adding language
            try:
                driver.find_elements_by_xpath('//div[@class="select-label tail-select-container"]')[0].click()
                # sleep(3)
                lan_opts = driver.find_elements_by_xpath('//ul[@class="dropdown-optgroup"]')[2].find_elements_by_tag_name(
                    'li')
                for lan_opt in lan_opts:
                    if 'english' in lan_opt.text.lower():
                        lan_opt.click()
            except:
                try:
                    driver.find_elements_by_xpath('//div[@class="select-label tail-select-container"]')[0].click()
                    # sleep(3)
                    lan_opts = driver.find_elements_by_xpath('//ul[@class="dropdown-optgroup"]')[
                        2].find_elements_by_tag_name('li')
                    for lan_opt in lan_opts:
                        if 'english' in lan_opt.text.lower():
                            lan_opt.click()
                except:
                    error += 1
                    sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 11).value += ", Error while choosing language"
                    continue


            # choosing category:
            topic = sheet1.cell(row, 2).value
            try:
                driver.find_elements_by_xpath('//div[@class="select-label"]')[1].click()

                cat_opts = driver.find_elements_by_xpath('//ul[@class="dropdown-optgroup"]')[
                    1].find_elements_by_tag_name('li')
                for cat_opt in cat_opts:
                    # # # print(cat_opt.text, topic)
                    if topic.lower() in cat_opt.text.lower():
                        cat_opt.click()
            except:
                try:
                    driver.find_elements_by_xpath('//div[@class="select-label"]')[1].click()

                    cat_opts = driver.find_elements_by_xpath('//ul[@class="dropdown-optgroup"]')[
                        1].find_elements_by_tag_name('li')
                    for cat_opt in cat_opts:
                        # # # print(cat_opt.text, topic)
                        if topic.lower() in cat_opt.text.lower():
                            cat_opt.click()

                except:
                    error += 1
                    sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 11).value += ", Error while choosing category"
                    sheet1.cell(row, 2).fill = PatternFill('solid', fgColor='FF0000')
                    continue

            # adding title:
            try:
                title = sheet1.cell(row, 3).value
                js = "document.getElementsByClassName('input w-full border mt-2')[0].value = '%s';"%title
                driver.execute_script(js)
                driver.find_element_by_xpath('//input[@name="title"]').click()


            except:
                try:
                    title = sheet1.cell(row, 3).value
                    driver.find_element_by_xpath('//input[@name="title"]').send_keys(title)
                    driver.find_element_by_xpath('//input[@name="title"]').click()
                except:
                    try:
                        title = sheet1.cell(row, 3).value
                        driver.find_element_by_xpath('//input[@name="title"]').click()

                        past = pyperclip.paste()
                        pyperclip.copy(title)
                        ActionChains(driver).send_keys(Keys.CONTROL, 'v')
                        pyperclip.copy(past)
                        driver.find_element_by_xpath('//input[@name="title"]').click()
                    except:
                        sheet1.cell(row, 3).fill = PatternFill('solid', fgColor='FF0000')
                        error += 1
                        sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 11).value += ", Error while entering title"
                        continue


            # adding discription
            try:
                discription = sheet1.cell(row, 4).value
                driver.find_element_by_xpath('//iframe[@class="cke_wysiwyg_frame cke_reset"]').click()

                past = pyperclip.paste()
                pyperclip.copy(discription)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('v').key_up(Keys.CONTROL).perform()
                pyperclip.copy(past)

            except:
                try:
                    discription = sheet1.cell(row, 4).value
                    driver.find_element_by_xpath('//iframe[@class="cke_wysiwyg_frame cke_reset"]').send_keys(discription)
                except:
                    try:
                        discription = sheet1.cell(row, 4).value
                        driver.find_element_by_xpath('//iframe[@class="cke_wysiwyg_frame cke_reset"]').click()

                        past = pyperclip.paste()
                        pyperclip.copy(discription)
                        ActionChains(driver).key_down(Keys.CONTROL).send_keys('v').key_up(Keys.CONTROL).perform()
                        pyperclip.copy(past)
                    except:
                        sheet1.cell(row, 4).fill = PatternFill('solid', fgColor='FF0000')
                        error += 1
                        sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 11).value += ", Error while entering description"
                        continue
            # adding discription again for 1st blog
            if row == 2:
                time.sleep(1)
                try:
                    discription = sheet1.cell(row, 4).value
                    driver.find_element_by_xpath('//iframe[@class="cke_wysiwyg_frame cke_reset"]').clear()
                    driver.find_element_by_xpath('//iframe[@class="cke_wysiwyg_frame cke_reset"]').send_keys(discription)
                except:
                    try:
                        discription = sheet1.cell(row, 4).value
                        driver.find_element_by_xpath('//iframe[@class="cke_wysiwyg_frame cke_reset"]').click()

                        past = pyperclip.paste()
                        pyperclip.copy(discription)
                        ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
                        ActionChains(driver).key_down(Keys.CONTROL).send_keys('v').key_up(Keys.CONTROL).perform()
                        pyperclip.copy(past)
                    except:
                        sheet1.cell(row, 4).fill = PatternFill('solid', fgColor='FF0000')
                        error += 1
                        sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 11).value += ", Error while entering description"
                        continue

            # adding url
            try:
                url = sheet1.cell(row, 5).value
                js = "document.getElementsByClassName('input w-full border mt-2')[3].value = '%s';"%url
                driver.execute_script(js)

            except:
                try:
                    url = sheet1.cell(row, 5).value
                    driver.find_element_by_xpath('//input[@name="url"]').send_keys(url)
                except:
                    try:
                        url = sheet1.cell(row, 5).value
                        driver.find_element_by_xpath('//input[@name="url"]').click()

                        past = pyperclip.paste()
                        pyperclip.copy(url)
                        ActionChains(driver).send_keys(Keys.CONTROL, 'v')
                        pyperclip.copy(past)
                    except:
                        sheet1.cell(row, 5).fill = PatternFill('solid', fgColor='FF0000')
                        error += 1
                        sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 11).value += ", Error while entering url"
                        continue

            # scheduling
            try:
                currentTime = time.strftime("%H:%M")
                js = 'document.getElementsByName("schedule_time")[0].value = "%s";' % str(currentTime)
                driver.execute_script(js)
            except:
                try:
                    currentTime = time.strftime("%H:%M")
                    js = 'document.getElementsByName("schedule_time")[0].value = "%s";' % str(currentTime)
                    driver.execute_script(js)
                except:
                    error += 1
                    sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 11).value += ", Error while choosing time"
                    continue

            # clicking visibility options
            try:
                for opt in driver.find_elements_by_xpath('//input[@class="input border mr-2"]'):
                    opt.click()
            except:
                try:
                    for opt in driver.find_elements_by_xpath('//input[@class="input border mr-2"]'):
                        opt.click()
                except:
                    error += 1
                    sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 11).value += ", Error while checking visibility options"
                    continue


            # clicking create:
            try:
                time.sleep(2)
                js = "document.getElementById('createBtn').click();"
                driver.execute_script(js)
                # try:
                #     time.sleep(2)
                #     js = "document.getElementById('createBtn').click();"
                #     driver.execute_script(js)
                # except:
                #     try:
                #         driver.find_element_by_id('createBtn').click()
                #     except:
                #         pass
            except:
                try:
                    driver.find_element_by_id('createBtn').click()
                except:
                    error += 1
                    sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 11).value += ", Error while clicking create btn"
                    continue

            # tracking update:
            try:
                driver.find_element_by_xpath('//table[@class="table table-report -mt-2"]')
                sent += 1
                sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='00FF00')
                sheet1.cell(row, 11).value += ", uploaded"
            except:
                # time.sleep(30)
                try:
                    time.sleep(2)
                    js = "document.getElementById('createBtn').click();"
                    driver.execute_script(js)
                    # try:
                    #     time.sleep(2)
                    #     js = "document.getElementById('createBtn').click();"
                    #     driver.execute_script(js)
                    # except:
                    #     try:
                    #         driver.find_element_by_id('createBtn').click()
                    #     except:
                    #         pass
                except:
                    try:
                        driver.find_element_by_id('createBtn').click()
                    except:
                        error += 1
                        sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 11).value += ", Error while clicking create btn"

                try:
                    driver.find_element_by_xpath('//table[@class="table table-report -mt-2"]')
                    sent += 1
                    sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='00FF00')
                    sheet1.cell(row, 11).value += ", uploaded"
                except:
                    sheet1.cell(row, 11).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 11).value += ", couldn't find blog after uploading"
                    continue
            # time.sleep(10)

        wb1.save('report.xlsx')



    if is_hin_var.get():
        for row in range(2,12):
            sheet1.cell(row, 12).value = ''
            try:
                driver.get('https://currentaffairsnews.in/add-blog/side-menu/light')
            except:
                driver.get('https://currentaffairsnews.in/add-blog/side-menu/light')

            # img uploading
            img = sheet1.cell(row, 10).value

            try:
                driver.find_elements_by_xpath('//input[@id="image"]')[0].send_keys(img)
            except:
                try:
                    driver.find_elements_by_xpath('//input[@id="image"]')[0].send_keys(img)
                except:
                    error += 1
                    sheet1.cell(row, 10).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 12).value += ", Error while uploading img"
                    continue


            # adding language
            try:
                driver.find_elements_by_xpath('//div[@class="select-label tail-select-container"]')[0].click()
                # sleep(3)
                lan_opts = driver.find_elements_by_xpath('//ul[@class="dropdown-optgroup"]')[2].find_elements_by_tag_name(
                    'li')
                for lan_opt in lan_opts:
                    if 'hindi' in lan_opt.text.lower():
                        lan_opt.click()
            except:
                try:
                    driver.find_elements_by_xpath('//div[@class="select-label tail-select-container"]')[0].click()
                    # sleep(3)
                    lan_opts = driver.find_elements_by_xpath('//ul[@class="dropdown-optgroup"]')[
                        2].find_elements_by_tag_name('li')
                    for lan_opt in lan_opts:
                        if 'hindi' in lan_opt.text.lower():
                            lan_opt.click()
                except:
                    error += 1
                    sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 12).value += ", Error while choosing language"
                    continue


            # choosing category:
            topic = sheet1.cell(row, 2).value
            try:
                driver.find_elements_by_xpath('//div[@class="select-label"]')[1].click()

                cat_opts = driver.find_elements_by_xpath('//ul[@class="dropdown-optgroup"]')[
                    1].find_elements_by_tag_name('li')
                for cat_opt in cat_opts:
                    # # # print(cat_opt.text, topic)
                    if topic.lower() in cat_opt.text.lower():
                        cat_opt.click()
            except:
                try:
                    driver.find_elements_by_xpath('//div[@class="select-label"]')[1].click()

                    cat_opts = driver.find_elements_by_xpath('//ul[@class="dropdown-optgroup"]')[
                        1].find_elements_by_tag_name('li')
                    for cat_opt in cat_opts:
                        # # # print(cat_opt.text, topic)
                        if topic.lower() in cat_opt.text.lower():
                            cat_opt.click()

                except:
                    error += 1
                    sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 12).value += ", Error while choosing category"
                    sheet1.cell(row, 2).fill = PatternFill('solid', fgColor='FF0000')
                    continue

            # adding title:
            try:
                title = sheet1.cell(row, 6).value
                js = "document.getElementsByClassName('input w-full border mt-2')[0].value = '%s';"%title
                driver.execute_script(js)
                # driver.find_element_by_xpath('//input[@name="title"]').click()


            except:
                try:
                    title = sheet1.cell(row, 6).value
                    driver.find_element_by_xpath('//input[@name="title"]').send_keys(title)
                except:
                    try:
                        title = sheet1.cell(row, 6).value
                        driver.find_element_by_xpath('//input[@name="title"]').click()

                        past = pyperclip.paste()
                        pyperclip.copy(title)
                        ActionChains(driver).send_keys(Keys.CONTROL, 'v')
                        pyperclip.copy(past)
                    except:
                        sheet1.cell(row, 6).fill = PatternFill('solid', fgColor='FF0000')
                        error += 1
                        sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 12).value += ", Error while entering title"
                        continue

            # adding slug:
            try:
                slug = sheet1.cell(row, 7).value
                js = "document.getElementsByClassName('input w-full border mt-2')[1].value = '%s';" % slug
                driver.execute_script(js)
                # driver.find_element_by_xpath('//input[@name="title"]').click()


            except:
                try:
                    slug = sheet1.cell(row, 7).value
                    driver.find_element_by_xpath('//input[@name="slug"]').clear()
                    driver.find_element_by_xpath('//input[@name="slug"]').send_keys(slug)
                except:
                    try:
                        slug = sheet1.cell(row, 7).value
                        driver.find_element_by_xpath('//input[@name="slug"]').click()
                        ActionChains(driver).send_keys(Keys.CONTROL, 'a')

                        past = pyperclip.paste()
                        pyperclip.copy(slug)
                        ActionChains(driver).send_keys(Keys.CONTROL, 'v')
                        pyperclip.copy(past)
                    except:
                        sheet1.cell(row, 7).fill = PatternFill('solid', fgColor='FF0000')
                        error += 1
                        sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 12).value += ", Error while entering title"
                        continue
            # if row == 7 or row == 8:
            #     time.sleep(0.5)
            #     try:
            #         slug = sheet1.cell(row, 7).value
            #         js = "document.getElementsByClassName('input w-full border mt-2')[1].value = '%s';" % slug
            #         driver.execute_script(js)
            #         # driver.find_element_by_xpath('//input[@name="title"]').click()
            #
            #
            #     except:
            #         try:
            #             slug = sheet1.cell(row, 7).value
            #             driver.find_element_by_xpath('//input[@name="slug"]').clear()
            #             driver.find_element_by_xpath('//input[@name="slug"]').send_keys(slug)
            #         except:
            #             try:
            #                 slug = sheet1.cell(row, 7).value
            #                 driver.find_element_by_xpath('//input[@name="slug"]').click()
            #                 ActionChains(driver).send_keys(Keys.CONTROL, 'a')
            #
            #                 past = pyperclip.paste()
            #                 pyperclip.copy(slug)
            #                 ActionChains(driver).send_keys(Keys.CONTROL, 'v')
            #                 pyperclip.copy(past)
            #             except:
            #                 sheet1.cell(row, 7).fill = PatternFill('solid', fgColor='FF0000')
            #                 error += 1
            #                 sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
            #                 sheet1.cell(row, 12).value += ", Error while entering title"
            #                 continue


            # adding discription
            try:
                discription = sheet1.cell(row, 8).value
                driver.find_element_by_xpath('//iframe[@class="cke_wysiwyg_frame cke_reset"]').click()

                past = pyperclip.paste()
                pyperclip.copy(discription)
                ActionChains(driver).key_down(Keys.CONTROL).send_keys('v').key_up(Keys.CONTROL).perform()
                pyperclip.copy(past)

            except:
                try:
                    discription = sheet1.cell(row, 8).value
                    driver.find_element_by_xpath('//iframe[@class="cke_wysiwyg_frame cke_reset"]').send_keys(discription)
                except:
                    try:
                        discription = sheet1.cell(row, 8).value
                        driver.find_element_by_xpath('//iframe[@class="cke_wysiwyg_frame cke_reset"]').click()

                        past = pyperclip.paste()
                        pyperclip.copy(discription)
                        ActionChains(driver).key_down(Keys.CONTROL).send_keys('v').key_up(Keys.CONTROL).perform()
                        pyperclip.copy(past)
                    except:
                        sheet1.cell(row, 8).fill = PatternFill('solid', fgColor='FF0000')
                        error += 1
                        sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 12).value += ", Error while entering description"
                        continue


            # adding url
            try:
                url = sheet1.cell(row, 9).value
                js = "document.getElementsByClassName('input w-full border mt-2')[3].value = '%s';"%url
                driver.execute_script(js)

            except:
                try:
                    url = sheet1.cell(row, 9).value
                    driver.find_element_by_xpath('//input[@name="url"]').send_keys(url)
                except:
                    try:
                        url = sheet1.cell(row, 9).value
                        driver.find_element_by_xpath('//input[@name="url"]').click()

                        past = pyperclip.paste()
                        pyperclip.copy(url)
                        ActionChains(driver).send_keys(Keys.CONTROL, 'v')
                        pyperclip.copy(past)
                    except:
                        sheet1.cell(row, 9).fill = PatternFill('solid', fgColor='FF0000')
                        error += 1
                        sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 12).value += ", Error while entering url"
                        continue

            # scheduling
            try:
                schedule_time = (datetime.now() + timedelta(hours=1)).strftime('%H:%M')
                js = 'document.getElementsByName("schedule_time")[0].value = "%s";' % str(schedule_time)
                driver.execute_script(js)
            except:
                try:
                    schedule_time = (datetime.now() + timedelta(hours=1)).strftime('%H:%M')
                    js = 'document.getElementsByName("schedule_time")[0].value = "%s";' % str(schedule_time)
                    driver.execute_script(js)
                except:
                    error += 1
                    sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 12).value += ", Error while choosing time"
                    continue

            # clicking visibility options
            try:
                for opt in driver.find_elements_by_xpath('//input[@class="input border mr-2"]'):
                    opt.click()
            except:
                try:
                    for opt in driver.find_elements_by_xpath('//input[@class="input border mr-2"]'):
                        opt.click()
                except:
                    error += 1
                    sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 12).value += ", Error while checking visibility options"
                    continue


            # clicking create:
            try:
                time.sleep(1)
                js = "document.getElementById('createBtn').click();"
                driver.execute_script(js)
                # try:
                #     time.sleep(2)
                #     js = "document.getElementById('createBtn').click();"
                #     driver.execute_script(js)
                # except:
                #     try:
                #         driver.find_element_by_id('createBtn').click()
                #     except:
                #         pass
            except:
                try:
                    driver.find_element_by_id('createBtn').click()
                except:
                    error += 1
                    sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 12).value += ", Error while clicking create btn"
                    continue

            # tracking update:
            try:
                driver.find_element_by_xpath('//table[@class="table table-report -mt-2"]')
                sent += 1
                sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='00FF00')
                sheet1.cell(row, 12).value += ", uploaded"
            except:
                # trying to add slug again :
                try:
                    slug = sheet1.cell(row, 7).value
                    js = "document.getElementsByClassName('input w-full border mt-2')[1].value = '%s';" % slug
                    driver.execute_script(js)
                    # driver.find_element_by_xpath('//input[@name="title"]').click()


                except:
                    try:
                        slug = sheet1.cell(row, 7).value
                        driver.find_element_by_xpath('//input[@name="slug"]').clear()
                        driver.find_element_by_xpath('//input[@name="slug"]').send_keys(slug)
                    except:
                        try:
                            slug = sheet1.cell(row, 7).value
                            driver.find_element_by_xpath('//input[@name="slug"]').click()
                            ActionChains(driver).send_keys(Keys.CONTROL, 'a')

                            past = pyperclip.paste()
                            pyperclip.copy(slug)
                            ActionChains(driver).send_keys(Keys.CONTROL, 'v')
                            pyperclip.copy(past)
                        except:
                            sheet1.cell(row, 7).fill = PatternFill('solid', fgColor='FF0000')
                            error += 1
                            sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                            sheet1.cell(row, 12).value += ", Error while entering title"
                try:
                    time.sleep(1)
                    js = "document.getElementById('createBtn').click();"
                    driver.execute_script(js)
                    # try:
                    #     time.sleep(2)
                    #     js = "document.getElementById('createBtn').click();"
                    #     driver.execute_script(js)
                    # except:
                    #     try:
                    #         driver.find_element_by_id('createBtn').click()
                    #     except:
                    #         pass
                except:
                    try:
                        driver.find_element_by_id('createBtn').click()
                    except:
                        error += 1
                        sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                        sheet1.cell(row, 12).value += ", Error while clicking create btn"

                try:
                    driver.find_element_by_xpath('//table[@class="table table-report -mt-2"]')
                    sent += 1
                    sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='00FF00')
                    sheet1.cell(row, 12).value += ", uploaded"
                except:
                    sheet1.cell(row, 12).fill = PatternFill('solid', fgColor='FF0000')
                    sheet1.cell(row, 12).value += ", couldn't find blog after uploading"
                    continue
            # time.sleep(10)


        wb1.save('report.xlsx')
    # time.sleep(2)
    driver.quit()
    popup('Uploading complete\nTotal Error: %s\nTotal Blog Uploaded: %s'%(error, sent), 'report.xlsx', 'notification')
    return

upload_btn = tk.Button(main_window, text='Upload',font=('Lucida Calligraphy', 10, 'bold'),bg = "#000000", fg = '#FF00D8', command=upload)
upload_btn.place(relx=0.45, rely=0.7, relwidth=0.09, relheight=0.05)

with open('showBrowser.txt', 'r', encoding="utf-8") as file:
    if int(file.readline()):
        is_showBrowser_button.select()
        showBrowser()
    else:
        pass


with open('autoDetector.txt', 'r', encoding="utf-8") as file:
    if int(file.readline()):
        is_autoDetector_button.select()
        autoDetector()
    else:
        pass

with open('autoAnalyser.txt', 'r', encoding="utf-8") as file:
    if int(file.readline()):
        is_autoAnalyser_button.select()
        autoAnalyser()
    else:
        pass

with open('autoUploader.txt', 'r', encoding="utf-8") as file:
    if int(file.readline()):
        is_autoUploader_button.select()
        autoUploader()
    else:
        pass

main_window.mainloop()